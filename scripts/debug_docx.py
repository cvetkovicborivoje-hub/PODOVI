#!/usr/bin/env python3
"""Debug script to see what's actually in the DOCX files"""

from docx import Document
from pathlib import Path

# Read the first DOCX file
docx_file = Path("public/documents/lvt/creation-30/63876 - CREATION 30 - Product description.docx")

print("=" * 80)
print(f"Reading: {docx_file.name}")
print("=" * 80)

doc = Document(docx_file)

print("\n[PARAGRAPHS]")
print("-" * 80)
for i, para in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n[TABLES]")
print("-" * 80)
for table_idx, table in enumerate(doc.tables):
    print(f"\nTable {table_idx}:")
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {' | '.join(cells)}")
