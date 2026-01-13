#!/usr/bin/env python3
"""
Extract specifications from LVT product documentation
Reads PDF and DOCX files from public/documents/lvt/ folders
"""

import os
import json
import re
from pathlib import Path
import PyPDF2
from docx import Document
import pdfplumber

def read_pdf_text(pdf_path):
    """Extract text from PDF using pdfplumber (better for tables)"""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def read_docx_text(docx_path):
    """Extract text from DOCX"""
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text

def extract_specs_from_text(text, product_name):
    """Extract key specifications from text"""
    specs = {
        "product": product_name,
        "thickness": None,
        "wear_layer": None,
        "dimensions": None,
        "format": None,
        "usage_class": None,
        "colors": None,
        "fire_class": None,
        "acoustic": None,
        "installation": None
    }
    
    # Patterns for extraction - adapted for Gerflor format
    patterns = {
        "thickness": r"total thickness of (\d+[.,]?\d*\s*mm)",
        "wear_layer": r"(\d+[.,]\d+mm)\s+wear[\s-]layer",
        "dimensions": r"(\d+\s*[x×]\s*\d+\s*mm)",
        "usage_class": r"(?:classification|usage class)\s+(\d+-\d+)",
        "colors": r"(\d+)\s+(?:colors|designs|references|decors)",
        "fire_class": r"(?:classified as|fire classification:?)\s+([A-F][a-z]*-?s?\d?)",
        "acoustic": r"(\d+)\s*dB",
    }
    
    text_lower = text.lower()
    
    # Extract using patterns
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            specs[key] = match.group(1).strip()
    
    # Determine format
    if "plank" in text_lower:
        specs["format"] = "plank"
    elif "tile" in text_lower:
        specs["format"] = "tile"
    
    # Determine installation type
    if "clic" in product_name.lower() or "click" in product_name.lower():
        specs["installation"] = "click system"
    elif "looselay" in product_name.lower():
        specs["installation"] = "loose lay"
    elif "connect" in product_name.lower():
        specs["installation"] = "connect system"
    else:
        specs["installation"] = "glue down"
    
    # Check for acoustic
    if "acoustic" in product_name.lower():
        specs["acoustic"] = "Yes"
    
    return specs

def scan_lvt_documents():
    """Scan all LVT folders and extract specifications"""
    base_path = Path("public/documents/lvt")
    all_specs = {}
    
    # Iterate through all collection folders
    for folder in base_path.iterdir():
        if not folder.is_dir() or folder.name == "general":
            continue
        
        print(f"\n[*] Scanning: {folder.name}")
        product_name = folder.name.replace("-", " ").title()
        
        # Look for product description or technical data sheet files
        description_files = []
        for file in folder.iterdir():
            fname_lower = file.name.lower().replace("-", " ").replace("_", " ")
            if ("product description" in fname_lower or 
                "technical data sheet" in fname_lower or
                "technical datasheet" in fname_lower or
                "tech datasheet" in fname_lower):
                description_files.append(file)
        
        if not description_files:
            print(f"  [!] No product description found")
            continue
        
        # Sort to prefer PDF files (more reliable)
        pdf_files = [f for f in description_files if f.suffix.lower() == ".pdf"]
        docx_files = [f for f in description_files if f.suffix.lower() in [".docx", ".doc"]]
        
        # Try PDF first, then DOCX
        desc_file = pdf_files[0] if pdf_files else (docx_files[0] if docx_files else description_files[0])
        print(f"  [+] Reading: {desc_file.name}")
        
        text = ""
        if desc_file.suffix.lower() == ".pdf":
            text = read_pdf_text(desc_file)
        elif desc_file.suffix.lower() in [".docx", ".doc"]:
            text = read_docx_text(desc_file)
        
        if text:
            print(f"  [OK] Extracted {len(text)} characters")
            specs = extract_specs_from_text(text, folder.name)
            all_specs[folder.name] = specs
            
            # Print found specs
            print(f"  [SPECS]:")
            for key, value in specs.items():
                if value and key != "product":
                    print(f"     - {key}: {value}")
        else:
            print(f"  [ERROR] Could not extract text")
    
    return all_specs

def main():
    print("=" * 60)
    print("LVT SPECIFICATION EXTRACTOR")
    print("=" * 60)
    
    specs = scan_lvt_documents()
    
    # Save to JSON
    output_file = "scripts/lvt_specs_extracted.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(specs, f, indent=2, ensure_ascii=False)
    
    print("\n" + "=" * 60)
    print(f"DONE! Extracted specs for {len(specs)} products")
    print(f"Saved to: {output_file}")
    print("=" * 60)

if __name__ == "__main__":
    main()
