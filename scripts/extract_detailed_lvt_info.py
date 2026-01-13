#!/usr/bin/env python3
"""
Extract DETAILED information from LVT documents
Focus on: colors, dimensions, warranty, certifications, applications, maintenance
"""

import json
import re
from pathlib import Path
from docx import Document
import pdfplumber

def read_pdf_text(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def read_docx_text(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text

def extract_detailed_info(text, folder_name):
    """Extract detailed product information"""
    info = {
        "folder": folder_name,
        "colors_count": None,
        "dimensions": [],
        "plank_sizes": [],
        "tile_sizes": [],
        "warranty_years": None,
        "certifications": [],
        "eco_features": [],
        "applications": [],
        "maintenance": None,
        "recycled_content": None,
        "phthalate_free": False,
        "voc_emission": None,
        "antimicrobial": False,
        "underfloor_heating": False,
    }
    
    text_lower = text.lower()
    
    # Colors/Designs/References
    color_patterns = [
        r'(\d+)\s+(?:colors|colours|designs|decors|references)',
        r'available in (\d+)',
        r'(\d+)\s+different'
    ]
    for pattern in color_patterns:
        match = re.search(pattern, text_lower)
        if match:
            info["colors_count"] = int(match.group(1))
            break
    
    # Dimensions - planks
    plank_patterns = [
        r'(\d+(?:\.\d+)?)\s*[x×]\s*(\d+(?:\.\d+)?)\s*mm',
        r'(\d+)\s*mm\s*[x×]\s*(\d+)\s*mm',
        r'dimensions?[:\s]+(\d+)\s*[x×]\s*(\d+)'
    ]
    for pattern in plank_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            width, length = match
            size = f"{length} x {width} mm"
            if size not in info["dimensions"]:
                info["dimensions"].append(size)
                # Classify as plank or tile based on size
                if int(float(length)) > 600:  # Usually planks are longer
                    info["plank_sizes"].append(size)
                else:
                    info["tile_sizes"].append(size)
    
    # Warranty
    warranty_patterns = [
        r'(\d+)[\s-]years?\s+warranty',
        r'warranty[:\s]+(\d+)\s+years?',
        r'guarantee[:\s]+(\d+)\s+years?'
    ]
    for pattern in warranty_patterns:
        match = re.search(pattern, text_lower)
        if match:
            info["warranty_years"] = int(match.group(1))
            break
    
    # Certifications
    certs = {
        "FloorScore": r'floorscore',
        "Indoor Air Comfort Gold": r'indoor air comfort gold|iac gold',
        "M1": r'\bm1\b',
        "A+": r'a\+\s+(?:rating|label)',
        "CE": r'\bce\b',
        "ISO 9001": r'iso\s*9001',
        "ISO 14001": r'iso\s*14001',
        "REACH": r'reach',
        "EPD": r'\bepd\b',
        "Cradle to Cradle": r'cradle to cradle|c2c'
    }
    for cert_name, pattern in certs.items():
        if re.search(pattern, text_lower):
            info["certifications"].append(cert_name)
    
    # Eco features
    if re.search(r'phthalate[\s-]free|without phthalates|free from phthalates', text_lower):
        info["phthalate_free"] = True
        info["eco_features"].append("Bez ftalata")
    
    if re.search(r'100%\s+recyclable|fully recyclable', text_lower):
        info["eco_features"].append("100% reciklabilno")
    
    recycled_match = re.search(r'(\d+)%\s+recycled content', text_lower)
    if recycled_match:
        info["recycled_content"] = int(recycled_match.group(1))
        info["eco_features"].append(f"{recycled_match.group(1)}% recikliranog sadržaja")
    
    voc_match = re.search(r'tvoc.*?below\s+(\d+)', text_lower)
    if voc_match:
        info["voc_emission"] = f"< {voc_match.group(1)} µg/m³"
        info["eco_features"].append("Niske VOC emisije")
    
    # Applications
    if re.search(r'residential|domestic', text_lower):
        info["applications"].append("Stambeni prostori")
    if re.search(r'commercial|public', text_lower):
        info["applications"].append("Komercijalni prostori")
    if re.search(r'heavy commercial|industrial', text_lower):
        info["applications"].append("Teška komercijalna upotreba")
    
    # Antimicrobial
    if re.search(r'antimicrobial|antibacterial|bacteriostatic', text_lower):
        info["antimicrobial"] = True
    
    # Underfloor heating
    if re.search(r'underfloor heating|floor heating compatible', text_lower):
        info["underfloor_heating"] = True
    
    # Maintenance
    if re.search(r'easy to clean|easy maintenance', text_lower):
        info["maintenance"] = "Lako održavanje"
    
    return info

def scan_all_lvt_docs():
    """Scan all LVT documentation folders"""
    base_path = Path("public/documents/lvt")
    all_info = {}
    
    for folder in base_path.iterdir():
        if not folder.is_dir() or folder.name == "general":
            continue
        
        print(f"\n[*] Scanning: {folder.name}")
        
        # Collect all text from all files in folder
        all_text = ""
        file_count = 0
        
        for file in folder.iterdir():
            if file.suffix.lower() == '.pdf':
                text = read_pdf_text(file)
                if text:
                    all_text += text + "\n\n"
                    file_count += 1
            elif file.suffix.lower() in ['.docx', '.doc']:
                text = read_docx_text(file)
                if text:
                    all_text += text + "\n\n"
                    file_count += 1
        
        print(f"  [+] Read {file_count} files, {len(all_text)} characters")
        
        if all_text:
            info = extract_detailed_info(all_text, folder.name)
            all_info[folder.name] = info
            
            # Print findings
            if info["colors_count"]:
                print(f"  Colors: {info['colors_count']}")
            if info["dimensions"]:
                print(f"  Dimensions: {', '.join(info['dimensions'][:3])}")
            if info["warranty_years"]:
                print(f"  Warranty: {info['warranty_years']} years")
            if info["certifications"]:
                print(f"  Certifications: {', '.join(info['certifications'][:3])}")
            if info["eco_features"]:
                print(f"  Eco: {', '.join(info['eco_features'][:2])}")
    
    return all_info

print("=" * 70)
print("DETAILED LVT INFORMATION EXTRACTION")
print("=" * 70)

info = scan_all_lvt_docs()

# Save to JSON
output_file = "scripts/lvt_detailed_info.json"
with open(output_file, 'w', encoding='utf-8') as f:
    json.dump(info, f, indent=2, ensure_ascii=False)

print("\n" + "=" * 70)
print(f"DONE! Detailed info for {len(info)} collections saved to {output_file}")
print("=" * 70)
